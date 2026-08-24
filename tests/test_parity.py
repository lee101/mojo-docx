from __future__ import annotations

from base64 import b64decode
from datetime import datetime
from io import BytesIO
import json
import zipfile

import pytest

from docx import Document
from docx._lib import PARALLEL_ESCAPE_BYTES, analyze_xml, escape_batch, lib
from docx.enum.section import WD_ORIENT, WD_SECTION_START
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.image import read_image
from docx.opc.package import OpcPackage
from docx.shared import Cm, Inches, Pt, RGBColor

PNG_1X1 = b64decode(
    "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mNk"
    "+A8AAQUBAScY42YAAAAASUVORK5CYII="
)
JPEG_3X2 = b"\xff\xd8\xff\xc0\x00\x11\x08\x00\x02\x00\x03" + b"\x00" * 10
GIF_3X2 = b"GIF89a\x03\x00\x02\x00" + b"\x00" * 8


def save_bytes(document) -> bytes:
    stream = BytesIO()
    document.save(stream)
    return stream.getvalue()


def test_new_document_defaults_match_upstream(upstream):
    reference = json.loads(
        upstream(
            """
            import json
            from docx import Document
            d = Document()
            s = d.sections[0]
            print(json.dumps([len(d.paragraphs), len(d.tables), len(d.sections),
                              s.page_width, s.page_height, s.top_margin,
                              d.core_properties.author, d.core_properties.comments]))
            """
        )
    )
    document = Document()
    section = document.sections[0]
    ours = [
        len(document.paragraphs),
        len(document.tables),
        len(document.sections),
        section.page_width,
        section.page_height,
        section.top_margin,
        document.core_properties.author,
        document.core_properties.comments,
    ]
    assert ours == reference


@pytest.mark.parametrize(
    "value",
    [
        "plain",
        " leading and trailing ",
        "A & B < C > D",
        "tabs\tand\nlines\r\nwork",
        "Zażółć gęślą jaźń — 東京",
        'quotes " stay literal in text',
    ],
)
def test_paragraph_text_roundtrip_matches_upstream(value, upstream, tmp_path):
    path = tmp_path / "ours.docx"
    document = Document()
    document.add_paragraph(value)
    document.save(path)
    summary = json.loads(
        upstream(
            """
            import json, sys
            from docx import Document
            d = Document(sys.argv[1])
            print(json.dumps([[p.text, [r.text for r in p.runs]] for p in d.paragraphs]))
            """,
            path,
        )
    )
    assert summary == [[value.replace("\r\n", "\n").replace("\r", "\n"), [
        value.replace("\r\n", "\n").replace("\r", "\n")
    ]]]


def test_open_upstream_document_with_runs_and_table(upstream, tmp_path):
    path = tmp_path / "upstream.docx"
    upstream(
        """
        import sys
        from docx import Document
        d = Document()
        d.add_heading("Heading", 2)
        p = d.add_paragraph("alpha")
        r = p.add_run("\\tbeta\\ngamma")
        r.bold = True
        t = d.add_table(2, 2)
        t.cell(0, 0).text = "A"
        t.cell(1, 1).text = "D"
        d.save(sys.argv[1])
        """,
        path,
    )
    document = Document(path)
    assert [paragraph.text for paragraph in document.paragraphs] == [
        "Heading",
        "alpha\tbeta\ngamma",
    ]
    assert document.paragraphs[0].style.name == "Heading 2"
    assert document.paragraphs[1].runs[1].bold is True
    assert [[cell.text for cell in row.cells] for row in document.tables[0].rows] == [
        ["A", ""],
        ["", "D"],
    ]


def test_heading_styles_match_upstream_names():
    document = Document()
    assert document.add_heading("Title", 0).style.name == "Title"
    for level in range(1, 10):
        assert document.add_heading(str(level), level).style.name == f"Heading {level}"
    with pytest.raises(ValueError):
        document.add_heading("bad", 10)


def test_run_formatting_is_readable_by_upstream(upstream, tmp_path):
    path = tmp_path / "formatting.docx"
    document = Document()
    run = document.add_paragraph().add_run("formatted")
    run.bold = True
    run.italic = False
    run.underline = True
    run.font.name = "Aptos"
    run.font.size = Pt(13.5)
    run.font.color.rgb = RGBColor(0x12, 0xAB, 0xF0)
    document.save(path)
    got = json.loads(
        upstream(
            """
            import json, sys
            from docx import Document
            r = Document(sys.argv[1]).paragraphs[0].runs[0]
            print(json.dumps([r.text, r.bold, r.italic, bool(r.underline),
                              r.font.name, r.font.size.pt, str(r.font.color.rgb)]))
            """,
            path,
        )
    )
    assert got == ["formatted", True, False, True, "Aptos", 13.5, "12ABF0"]

def test_run_character_style_roundtrip():
    run = Document().add_paragraph().add_run("styled", style="Emphasis")
    assert run.style.name == "Emphasis"
    reopened = Document(BytesIO(save_bytes(run._document))).paragraphs[0].runs[0]
    assert reopened.style.name == "Emphasis"


@pytest.mark.parametrize(
    ("alignment", "expected"),
    [
        (WD_ALIGN_PARAGRAPH.LEFT, 0),
        (WD_ALIGN_PARAGRAPH.CENTER, 1),
        (WD_ALIGN_PARAGRAPH.RIGHT, 2),
        (WD_ALIGN_PARAGRAPH.JUSTIFY, 3),
        (None, None),
    ],
)
def test_paragraph_alignment(alignment, expected):
    paragraph = Document().add_paragraph("x")
    paragraph.alignment = alignment
    assert paragraph.alignment == expected
    assert Document(BytesIO(save_bytes(paragraph._document))).paragraphs[0].alignment == expected


def test_paragraph_format_lengths_roundtrip():
    paragraph = Document().add_paragraph("x")
    formatting = paragraph.paragraph_format
    formatting.left_indent = Inches(0.5)
    formatting.right_indent = Cm(1)
    formatting.first_line_indent = Inches(-0.25)
    formatting.space_before = Pt(6)
    formatting.space_after = Pt(9)
    reopened = Document(BytesIO(save_bytes(paragraph._document))).paragraphs[0].paragraph_format
    assert reopened.left_indent.inches == pytest.approx(0.5, abs=0.001)
    assert reopened.right_indent.cm == pytest.approx(1, abs=0.01)
    assert reopened.first_line_indent.inches == pytest.approx(-0.25, abs=0.001)
    assert reopened.space_before.pt == pytest.approx(6)
    assert reopened.space_after.pt == pytest.approx(9)


def test_table_creation_and_growth():
    table = Document().add_table(2, 2, style="Table Grid")
    table.autofit = False
    table.cell(0, 0).text = "north"
    table.cell(1, 1).text = "south"
    row = table.add_row()
    row.cells[0].text = "third"
    column = table.add_column(Inches(1))
    column.cells[0].text = "east"
    assert len(table.rows) == 3
    assert len(table.columns) == 3
    assert table.style.name == "Table Grid"
    assert table.autofit is False
    assert table.cell(0, 2).text == "east"
    assert table.cell(2, 0).text == "third"

def test_nested_table_content_and_width_roundtrip():
    document = Document()
    outer = document.add_table(1, 1)
    outer.cell(0, 0).width = Inches(1.25)
    nested = outer.cell(0, 0).add_table(1, 1)
    nested.cell(0, 0).text = "nested"
    reopened = Document(BytesIO(save_bytes(document)))
    cell = reopened.tables[0].cell(0, 0)
    assert cell.width.inches == pytest.approx(1.25, abs=0.001)
    assert cell.tables[0].cell(0, 0).text == "nested"


def test_table_saved_here_matches_upstream_shape(upstream, tmp_path):
    path = tmp_path / "table.docx"
    document = Document()
    table = document.add_table(3, 2)
    for r, row in enumerate(table.rows):
        for c, cell in enumerate(row.cells):
            cell.text = f"{r}:{c}"
    document.save(path)
    got = json.loads(
        upstream(
            """
            import json, sys
            from docx import Document
            t = Document(sys.argv[1]).tables[0]
            print(json.dumps([[c.text for c in row.cells] for row in t.rows]))
            """,
            path,
        )
    )
    assert got == [[f"{r}:{c}" for c in range(2)] for r in range(3)]


def test_sections_and_margins():
    document = Document()
    section = document.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.left_margin = Inches(0.75)
    added = document.add_section(WD_SECTION_START.CONTINUOUS)
    assert len(document.sections) == 2
    assert document.sections[0].left_margin == Inches(0.75)
    assert added.orientation == WD_ORIENT.LANDSCAPE

def test_section_page_size_and_all_distances_roundtrip():
    document = Document()
    section = document.sections[0]
    values = {
        "page_width": Inches(8),
        "page_height": Inches(10),
        "top_margin": Inches(0.5),
        "bottom_margin": Inches(0.6),
        "left_margin": Inches(0.7),
        "right_margin": Inches(0.8),
        "header_distance": Inches(0.3),
        "footer_distance": Inches(0.4),
        "gutter": Inches(0.2),
    }
    for name, value in values.items():
        setattr(section, name, value)
    reopened = Document(BytesIO(save_bytes(document))).sections[0]
    for name, value in values.items():
        assert getattr(reopened, name) == value


def test_core_properties_roundtrip():
    document = Document()
    props = document.core_properties
    props.title = "Report"
    props.author = "Ada"
    props.keywords = "mojo, docx"
    props.created = datetime(2026, 7, 30, 12, 34, 56)
    props.revision = 7
    reopened = Document(BytesIO(save_bytes(document))).core_properties
    assert reopened.title == "Report"
    assert reopened.author == "Ada"
    assert reopened.keywords == "mojo, docx"
    assert reopened.created == datetime(2026, 7, 30, 12, 34, 56)
    assert reopened.revision == 7


def test_picture_dimensions_and_upstream_readability(upstream, tmp_path):
    path = tmp_path / "picture.docx"
    document = Document()
    shape = document.add_picture(BytesIO(PNG_1X1), width=Inches(2))
    document.save(path)
    assert shape.width == Inches(2)
    assert shape.height == Inches(2)
    got = json.loads(
        upstream(
            """
            import json, sys
            from docx import Document
            d = Document(sys.argv[1])
            s = d.inline_shapes[0]
            print(json.dumps([len(d.inline_shapes), s.width, s.height]))
            """,
            path,
        )
    )
    assert got == [1, Inches(2), Inches(2)]

@pytest.mark.parametrize(
    ("data", "extension", "width", "height"),
    [(PNG_1X1, "png", 1, 1), (JPEG_3X2, "jpg", 3, 2), (GIF_3X2, "gif", 3, 2)],
)
def test_supported_image_formats(data, extension, width, height):
    image = read_image(BytesIO(data))
    assert (image.extension, image.width_px, image.height_px) == (
        extension, width, height,
    )


def test_page_break_is_preserved_by_upstream(upstream, tmp_path):
    path = tmp_path / "break.docx"
    document = Document()
    document.add_page_break()
    document.save(path)
    result = upstream(
        """
        import sys
        from docx import Document
        from docx.oxml.ns import qn
        d = Document(sys.argv[1])
        print(d.paragraphs[0]._p.find('.//' + qn('w:br')).get(qn('w:type')))
        """,
        path,
    )
    assert result == "page"


def test_unknown_package_parts_are_preserved():
    document = Document()
    document._parts_data["customXml/item1.xml"] = b"<custom>payload</custom>"
    data = save_bytes(document)
    reopened = Document(BytesIO(data))
    assert reopened._parts_data["customXml/item1.xml"] == b"<custom>payload</custom>"
    assert "customXml/item1.xml" in zipfile.ZipFile(BytesIO(data)).namelist()


def test_opc_package_facade_saves():
    document = Document()
    document.add_paragraph("facade")
    stream = BytesIO()
    document.part.package.save(stream)
    stream.seek(0)
    assert Document(stream).paragraphs[0].text == "facade"
    assert any(part.partname == "/word/document.xml" for part in document.part.package.parts)


def test_mojo_escape_batch_matches_xml_rules():
    values = ["plain", "A&B<C>D", '"quoted"', "\tline\ncarriage\r", "東京"]
    assert escape_batch(values) == [
        "plain",
        "A&amp;B&lt;C&gt;D",
        '"quoted"',
        "\tline\ncarriage\r",
        "東京",
    ]
    assert escape_batch(values, attribute=True) == [
        "plain",
        "A&amp;B&lt;C&gt;D",
        "&quot;quoted&quot;",
        "&#9;line&#10;carriage&#13;",
        "東京",
    ]

def test_mojo_ffi_handles_empty_nul_and_rejects_wrong_types():
    assert escape_batch(["", "\0", "", "a\0b"]) == ["", "\0", "", "a\0b"]
    assert analyze_xml(b"") == {
        "elements": 0, "paragraphs": 0, "runs": 0,
        "texts": 0, "tables": 0, "rows": 0,
    }
    with pytest.raises(TypeError, match="list of strings"):
        escape_batch(("not", "a", "list"))
    with pytest.raises(TypeError, match="list of strings"):
        escape_batch(["valid", b"not text"])
    with pytest.raises(TypeError, match="must be bytes"):
        analyze_xml(bytearray(b"<w:p/>"))
    assert lib().mdx_escape_one(0, 0, 0, 0, 0, 0) == -1
    assert lib().mdx_escape_batch(0, 0, 0, 0, 0, 0, 0, 0, 0) == -1
    assert lib().mdx_analyze_xml(0, 0, 0) == -1


@pytest.mark.parametrize("length", [31, 32, 33, 63, 64, 65, 127])
def test_mojo_escape_simd_tail(length):
    value = "x" * (length - 1) + "&"
    assert escape_batch([value]) == ["x" * (length - 1) + "&amp;"]


def test_mojo_escape_multiple_entities_in_simd_block():
    value = '&<x>"\t\n\r' + "x" * 56
    assert escape_batch([value]) == ['&amp;&lt;x&gt;"\t\n\r' + "x" * 56]
    assert escape_batch([value], attribute=True) == [
        "&amp;&lt;x&gt;&quot;&#9;&#10;&#13;" + "x" * 56
    ]


@pytest.mark.parametrize(
    "length", [PARALLEL_ESCAPE_BYTES - 1, PARALLEL_ESCAPE_BYTES]
)
def test_mojo_escape_parallel_threshold(length):
    value = ("A<&>\t\"\n\r" * ((length // 8) + 1))[:length]
    expected = (
        value.replace("&", "&amp;")
        .replace("<", "&lt;")
        .replace(">", "&gt;")
        .replace('"', "&quot;")
        .replace("\t", "&#9;")
        .replace("\n", "&#10;")
        .replace("\r", "&#13;")
    )
    assert escape_batch([value], attribute=True) == [expected]


def test_loaded_document_cache_invalidates_after_each_mutation():
    original = Document()
    original.add_paragraph("before")
    reopened = Document(BytesIO(save_bytes(original)))

    first_cached = save_bytes(reopened)
    assert first_cached == save_bytes(reopened)

    reopened.paragraphs[0].text = "after"
    first_edit = save_bytes(reopened)
    assert Document(BytesIO(first_edit)).paragraphs[0].text == "after"

    reopened.core_properties.title = "changed after cached save"
    second_edit = save_bytes(reopened)
    assert Document(BytesIO(second_edit)).core_properties.title == "changed after cached save"


def test_mojo_xml_analyzer_counts_word_tags():
    document = Document()
    document.add_paragraph("one").add_run("two")
    document.add_table(2, 1)
    stats = analyze_xml(document.part.blob)
    assert stats["paragraphs"] == 3
    assert stats["runs"] == 2
    assert stats["texts"] == 2
    assert stats["tables"] == 1
    assert stats["rows"] == 2


@pytest.mark.parametrize("bad", [b"not a zip", b"PK\x03\x04broken"])
def test_invalid_packages_fail_cleanly(bad):
    with pytest.raises(ValueError, match="valid Word document package"):
        Document(BytesIO(bad))
