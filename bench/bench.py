"""Honest end-to-end benchmarks against python-docx on the same OOXML data."""

from __future__ import annotations

from io import BytesIO
import json
import math
import os
from pathlib import Path
import platform
import subprocess
import sys
import tempfile
import time


PARAGRAPHS = 20_000
TEXT = "A & B < C — WordprocessingML text with tabs\tand Unicode 東京. " * 2


def best_time(function, repeat: int = 3) -> float:
    best = math.inf
    for _ in range(repeat):
        start = time.perf_counter()
        function()
        best = min(best, time.perf_counter() - start)
    return best


def build_save(Document):
    def operation():
        document = Document()
        for index in range(PARAGRAPHS):
            document.add_paragraph(f"{index}: {TEXT}")
        stream = BytesIO()
        document.save(stream)
        return len(stream.getbuffer())

    return operation


def open_read(Document, path: str):
    def operation():
        document = Document(path)
        return sum(len(paragraph.text) for paragraph in document.paragraphs)

    return operation


def roundtrip_save(Document, path: str):
    document = Document(path)

    def operation():
        stream = BytesIO()
        document.save(stream)
        return len(stream.getbuffer())

    return operation


def table_save(Document):
    def operation():
        document = Document()
        table = document.add_table(200, 20)
        for row_index, row in enumerate(table.rows):
            for column_index, cell in enumerate(row.cells):
                cell.text = f"{row_index}:{column_index} & data"
        stream = BytesIO()
        document.save(stream)
        return len(stream.getbuffer())

    return operation


def upstream_worker(case: str, path: str) -> None:
    from docx import Document
    import docx

    if case == "build_save":
        operation = build_save(Document)
        repeat = 2
    elif case == "open_read":
        operation = open_read(Document, path)
        repeat = 3
    elif case == "roundtrip_save":
        operation = roundtrip_save(Document, path)
        repeat = 3
    elif case == "table_save":
        operation = table_save(Document)
        repeat = 2
    elif case == "xml_escape":
        from docx.oxml import OxmlElement
        from lxml import etree

        node = OxmlElement("w:t")
        node.text = TEXT * 65_536
        operation = lambda: etree.tostring(node)
        repeat = 3
    else:
        raise ValueError(case)
    operation()
    print(json.dumps({"seconds": best_time(operation, repeat), "version": docx.__version__}))


def upstream_time(case: str, path: Path) -> tuple[float, str]:
    environment = os.environ.copy()
    environment.pop("PYTHONPATH", None)
    process = subprocess.run(
        [sys.executable, str(Path(__file__).resolve()), "--upstream", case, str(path)],
        cwd=path.parent,
        env=environment,
        capture_output=True,
        text=True,
        check=True,
    )
    result = json.loads(process.stdout)
    return result["seconds"], result["version"]


def machine() -> str:
    cpu = platform.processor()
    if not cpu and Path("/proc/cpuinfo").exists():
        for line in Path("/proc/cpuinfo").read_text().splitlines():
            if line.startswith("model name"):
                cpu = line.split(":", 1)[1].strip()
                break
    return f"{cpu or 'unknown CPU'}, {platform.system()} {platform.release()}, Python {platform.python_version()}"


def main() -> None:
    from docx import Document
    from docx._lib import escape_batch

    with tempfile.TemporaryDirectory(prefix="mojo-docx-bench-") as directory:
        path = Path(directory) / "fixture.docx"
        fixture = Document()
        for index in range(PARAGRAPHS):
            fixture.add_paragraph(f"{index}: {TEXT}")
        fixture.save(path)

        escape_input = TEXT * 65_536
        cases = [
            ("Build + save 20k paragraphs", "build_save", build_save(Document), 2),
            ("Open + read 20k paragraphs", "open_read", open_read(Document, str(path)), 3),
            ("Round-trip save 20k paragraphs", "roundtrip_save", roundtrip_save(Document, str(path)), 3),
            ("Build + save 200x20 table", "table_save", table_save(Document), 2),
            ("Escape 8 MiB OOXML text", "xml_escape", lambda: escape_batch([escape_input]), 3),
        ]

        rows = []
        upstream_version = ""
        for label, key, ours, repeat in cases:
            ours()
            mojo_seconds = best_time(ours, repeat)
            reference_seconds, upstream_version = upstream_time(key, path)
            rows.append((label, mojo_seconds, reference_seconds))

    print(f"Machine: {machine()}")
    print()
    print(f"| workload | mojo-docx | python-docx {upstream_version} | upstream / Mojo |")
    print("| --- | ---: | ---: | ---: |")
    for label, mojo_seconds, reference_seconds in rows:
        print(
            f"| {label} | {mojo_seconds * 1000:.2f} ms | "
            f"{reference_seconds * 1000:.2f} ms | {reference_seconds / mojo_seconds:.2f}x |"
        )


if __name__ == "__main__":
    if len(sys.argv) >= 3 and sys.argv[1] == "--upstream":
        upstream_worker(sys.argv[2], sys.argv[3])
    else:
        main()
